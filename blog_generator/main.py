import sys
import os
import json
import google.generativeai as genai
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QLabel, QLineEdit, QTextEdit, QPushButton, 
                            QSlider, QMessageBox, QProgressBar, QComboBox, 
                            QStackedWidget, QFrame)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QPalette, QColor
from docx import Document
from datetime import datetime

# Worker threads remain the same as previous implementation
class GenerateTitlesWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, model, topic):
        super().__init__()
        self.model = model
        self.topic = topic

    def run(self):
        try:
            prompt = f"""Generate 10 unique, engaging, and SEO-friendly blog titles for the topic: {self.topic}
            Please provide numbered titles (1-10) that are creative and compelling."""
            
            response = self.model.generate_content(prompt)
            self.finished.emit(response.text)
        except Exception as e:
            self.error.emit(str(e))

class GenerateBlogWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, model, title, keywords, word_count):
        super().__init__()
        self.model = model
        self.title = title
        self.keywords = keywords
        self.word_count = word_count

    def run(self):
        try:
            prompt = f"""Write a comprehensive blog post with the following specifications:
            Title: {self.title}
            Keywords to include: {self.keywords}
            Target word count: {self.word_count}
            
            Please write an engaging, well-structured blog post that naturally incorporates the keywords
            and maintains a professional tone. Include an introduction, main body with subheadings,
            and a conclusion."""
            
            response = self.model.generate_content(prompt)
            self.finished.emit(response.text)
        except Exception as e:
            self.error.emit(str(e))

class APIKeyScreen(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAutoFillBackground(True)
        self.setup_ui()

    def setup_ui(self):
        # Set white background for this widget
        self.setStyleSheet("background-color: white;")
        
        layout = QVBoxLayout()
        
        # Welcome message
        welcome_label = QLabel("Welcome to Blog Generator")
        welcome_label.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        welcome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        welcome_label.setStyleSheet("color: black; background-color: white;")
        
        # API key input
        api_key_label = QLabel("Please enter your Google API Key:")
        api_key_label.setStyleSheet("color: black; background-color: white;")
        
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("Enter your API key here...")
        self.api_key_input.setMinimumWidth(300)
        self.api_key_input.setStyleSheet("""
            QLineEdit {
                color: black;
                background-color: white;
                border: 1px solid #ccc;
                border-radius: 4px;
                padding: 8px;
            }
        """)
        
        # Submit button
        self.submit_btn = QPushButton("Submit")
        self.submit_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                min-height: 30px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        
        # Add widgets to layout with some spacing
        layout.addStretch()
        layout.addWidget(welcome_label)
        layout.addSpacing(20)
        layout.addWidget(api_key_label)
        layout.addWidget(self.api_key_input)
        layout.addWidget(self.submit_btn)
        layout.addStretch()
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.setLayout(layout)

class MainScreen(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()

    def setup_ui(self):
        main_layout = QHBoxLayout()
        
        # Column 1 (Left side)
        left_column = QFrame()
        left_column.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Raised)
        left_column.setStyleSheet("QFrame { background-color: #f5f5f5; border-radius: 10px; padding: 10px; }")
        left_layout = QVBoxLayout()
        
        # Model selection
        model_label = QLabel("Select Model:")
        self.model_combo = QComboBox()
        models = ["gemini-pro", "gemini-flash", "gemini-1.5-pro", "gemini-1.5-flash",
                 "gemini-1.5-pro-002", "gemini-1.5-flash-002", "gemini-2.0-flash-exp",
                 "gemini-1.5-flash-8b"]
        self.model_combo.addItems(models)
        
        # Title input
        title_label = QLabel("Enter Topic:")
        self.title_input = QLineEdit()
        self.title_input.setStyleSheet("QLineEdit { padding: 5px; border: 1px solid #ccc; border-radius: 4px; }")
        
        # Generate titles button
        self.generate_titles_btn = QPushButton("Generate Titles")
        self.generate_titles_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        
        # Titles output
        self.titles_progress = QProgressBar()
        self.titles_progress.setVisible(False)
        self.titles_output = QTextEdit()
        self.titles_output.setReadOnly(True)
        self.titles_output.setStyleSheet("QTextEdit { background-color: white; border: 1px solid #ccc; border-radius: 4px; }")
        
        left_layout.addWidget(model_label)
        left_layout.addWidget(self.model_combo)
        left_layout.addWidget(title_label)
        left_layout.addWidget(self.title_input)
        left_layout.addWidget(self.generate_titles_btn)
        left_layout.addWidget(self.titles_progress)
        left_layout.addWidget(self.titles_output)
        left_column.setLayout(left_layout)
        
        # Column 2 (Right side)
        right_column = QFrame()
        right_column.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Raised)
        right_column.setStyleSheet("QFrame { background-color: #f5f5f5; border-radius: 10px; padding: 10px; }")
        right_layout = QVBoxLayout()
        
        # Selected title
        selected_title_label = QLabel("Selected Title:")
        self.selected_title_input = QLineEdit()
        self.selected_title_input.setStyleSheet("QLineEdit { padding: 5px; border: 1px solid #ccc; border-radius: 4px; }")
        
        # Keywords
        keywords_label = QLabel("Keywords (comma-separated):")
        self.keywords_input = QLineEdit()
        self.keywords_input.setStyleSheet("QLineEdit { padding: 5px; border: 1px solid #ccc; border-radius: 4px; }")
        
        # Word count slider
        word_count_label = QLabel("Word Count:")
        self.word_count_label = QLabel("500 words")
        self.word_count_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.word_count_slider = QSlider(Qt.Orientation.Horizontal)
        self.word_count_slider.setMinimum(100)
        self.word_count_slider.setMaximum(1000)
        self.word_count_slider.setValue(500)
        self.word_count_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.word_count_slider.setTickInterval(100)
        
        # Generate blog button
        self.generate_blog_btn = QPushButton("Generate Blog")
        self.generate_blog_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF4081;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F50057;
            }
        """)
        
        # Blog output
        self.blog_progress = QProgressBar()
        self.blog_progress.setVisible(False)
        self.blog_output = QTextEdit()
        self.blog_output.setReadOnly(True)
        self.blog_output.setStyleSheet("QTextEdit { background-color: white; border: 1px solid #ccc; border-radius: 4px; }")
        
        # Download button
        self.download_btn = QPushButton("Generate Document")
        self.download_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        
        right_layout.addWidget(selected_title_label)
        right_layout.addWidget(self.selected_title_input)
        right_layout.addWidget(keywords_label)
        right_layout.addWidget(self.keywords_input)
        right_layout.addWidget(word_count_label)
        right_layout.addWidget(self.word_count_slider)
        right_layout.addWidget(self.word_count_label)
        right_layout.addWidget(self.generate_blog_btn)
        right_layout.addWidget(self.blog_progress)
        right_layout.addWidget(self.blog_output)
        right_layout.addWidget(self.download_btn)
        right_column.setLayout(right_layout)
        
        # Add columns to main layout
        main_layout.addWidget(left_column)
        main_layout.addWidget(right_column)
        self.setLayout(main_layout)

class BlogGeneratorApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Blog Content Generator")
        self.setMinimumSize(1200, 800)

        # Set the main window background to white
        self.setStyleSheet("QMainWindow { background-color: white; }")
        
        # Create stacked widget for multiple screens
        self.stacked_widget = QStackedWidget()
        self.setCentralWidget(self.stacked_widget)
        
        # Create screens
        self.api_screen = APIKeyScreen()
        self.main_screen = MainScreen()
        
        # Add screens to stacked widget
        self.stacked_widget.addWidget(self.api_screen)
        self.stacked_widget.addWidget(self.main_screen)
        
        # Connect signals
        self.api_screen.submit_btn.clicked.connect(self.handle_api_key)
        self.main_screen.generate_titles_btn.clicked.connect(self.generate_titles)
        self.main_screen.generate_blog_btn.clicked.connect(self.generate_blog)
        self.main_screen.download_btn.clicked.connect(self.save_to_word)
        self.main_screen.word_count_slider.valueChanged.connect(self.update_word_count_label)
        
        # Check for saved API key
        self.check_saved_api_key()

        # Apply consistent styles to both screens
        self.api_screen.setStyleSheet("""
            QWidget {
                background-color: white;
            }
            QLabel {
                color: black;
                font-size: 14px;
                background-color: white;
            }
            QLineEdit {
                color: black;
                background-color: white;
                border: 1px solid #ccc;
                border-radius: 4px;
                padding: 8px;
                font-size: 14px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-size: 14px;
                font-weight: bold;
                min-height: 30px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)

        self.main_screen.setStyleSheet("""
            QLabel { color: black; }
            QLineEdit { 
                color: black;
                background-color: white;
            }
            QPushButton { 
                color: black;
                background-color: #f0f0f0;
            }
            QTextEdit { 
                color: black;
                background-color: white;
            }
            QComboBox {
                color: black;
                background-color: white;
            }
            QComboBox QAbstractItemView {
                color: black;
                background-color: white;
                selection-background-color: #a6a6a6;
            }
        """)
        
    def check_saved_api_key(self):
        try:
            if os.path.exists('config.json'):
                with open('config.json', 'r') as f:
                    config = json.load(f)
                    api_key = config.get('api_key')
                    if api_key:
                        genai.configure(api_key=api_key)
                        self.stacked_widget.setCurrentWidget(self.main_screen)
                        return
        except Exception:
            pass
        self.stacked_widget.setCurrentWidget(self.api_screen)
        
    def handle_api_key(self):
        api_key = self.api_screen.api_key_input.text().strip()
        if not api_key:
            msg = QMessageBox.warning(self, "Warning", "Please enter an API key!")
            return
            
        try:
            genai.configure(api_key=api_key)
            # Save API key
            with open('config.json', 'w') as f:
                json.dump({'api_key': api_key}, f)
            self.stacked_widget.setCurrentWidget(self.main_screen)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Invalid API key: {str(e)}")
    
    def update_word_count_label(self, value):
        self.main_screen.word_count_label.setText(f"{value} words")
    
    def generate_titles(self):
        topic = self.main_screen.title_input.text().strip()
        if not topic:
            QMessageBox.warning(self, "Warning", "Please enter a topic!")
            return
        
        model = genai.GenerativeModel(self.main_screen.model_combo.currentText())
        
        # Disable button and show progress
        self.main_screen.generate_titles_btn.setEnabled(False)
        self.main_screen.titles_progress.setVisible(True)
        self.main_screen.titles_progress.setRange(0, 0)
        
        # Create and start worker thread
        self.titles_worker = GenerateTitlesWorker(model, topic)
        self.titles_worker.finished.connect(self.on_titles_generated)
        self.titles_worker.error.connect(self.on_error)
        self.titles_worker.start()

    def generate_blog(self):
        title = self.main_screen.selected_title_input.text().strip()
        keywords = self.main_screen.keywords_input.text().strip()
        word_count = self.main_screen.word_count_slider.value()
        
        if not all([title, keywords]):
            QMessageBox.warning(self, "Warning", "Please fill in all fields!")
            return
        
        model = genai.GenerativeModel(self.main_screen.model_combo.currentText())
        
        # Disable button and show progress
        self.main_screen.generate_blog_btn.setEnabled(False)
        self.main_screen.blog_progress.setVisible(True)
        self.main_screen.blog_progress.setRange(0, 0)
        
        # Create and start worker thread
        self.blog_worker = GenerateBlogWorker(model, title, keywords, word_count)
        self.blog_worker.finished.connect(self.on_blog_generated)
        self.blog_worker.error.connect(self.on_error)
        self.blog_worker.start()

    def on_titles_generated(self, text):
        self.main_screen.titles_output.setText(text)
        self.main_screen.generate_titles_btn.setEnabled(True)
        self.main_screen.titles_progress.setVisible(False)

    def on_blog_generated(self, text):
        self.main_screen.blog_output.setText(text)
        self.main_screen.generate_blog_btn.setEnabled(True)
        self.main_screen.blog_progress.setVisible(False)

    def on_error(self, error_message):
        QMessageBox.critical(self, "Error", f"An error occurred: {error_message}")
        # Re-enable buttons and hide progress bars
        self.main_screen.generate_titles_btn.setEnabled(True)
        self.main_screen.generate_blog_btn.setEnabled(True)
        self.main_screen.titles_progress.setVisible(False)
        self.main_screen.blog_progress.setVisible(False)

    def save_to_word(self):
        if not self.main_screen.blog_output.toPlainText():
            QMessageBox.warning(self, "Warning", "No content to save!")
            return
            
        try:
            # Create a new Document
            doc = Document()
            
            # Add title
            doc.add_heading(self.main_screen.selected_title_input.text(), 0)
            
            # Add metadata
            doc.add_paragraph(f"Generated using: {self.main_screen.model_combo.currentText()}")
            doc.add_paragraph(f"Keywords: {self.main_screen.keywords_input.text()}")
            doc.add_paragraph(f"Target word count: {self.main_screen.word_count_slider.value()}")
            
            # Add a horizontal line
            doc.add_paragraph("_" * 50)
            
            # Add content
            doc.add_paragraph(self.main_screen.blog_output.toPlainText())
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"blog_{timestamp}.docx"
            
            # Save the document
            doc.save(filename)
            
            QMessageBox.information(self, "Success", f"Blog saved as {filename}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save document: {str(e)}")

def main():
    app = QApplication(sys.argv)
    
    # Set application style
    app.setStyle("Fusion")
    
    # Set dark palette for better contrast
    # palette = QPalette()
    # palette.setColor(QPalette.ColorRole.Window, QColor(53, 53, 53))
    # palette.setColor(QPalette.ColorRole.WindowText, Qt.GlobalColor.white)
    # palette.setColor(QPalette.ColorRole.Base, QColor(25, 25, 25))
    # palette.setColor(QPalette.ColorRole.AlternateBase, QColor(53, 53, 53))
    # palette.setColor(QPalette.ColorRole.ToolTipBase, Qt.GlobalColor.white)
    # palette.setColor(QPalette.ColorRole.ToolTipText, Qt.GlobalColor.white)
    # palette.setColor(QPalette.ColorRole.Text, Qt.GlobalColor.white)
    # palette.setColor(QPalette.ColorRole.Button, QColor(53, 53, 53))
    # palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
    # palette.setColor(QPalette.ColorRole.BrightText, Qt.GlobalColor.red)
    # palette.setColor(QPalette.ColorRole.Link, QColor(42, 130, 218))
    # palette.setColor(QPalette.ColorRole.Highlight, QColor(42, 130, 218))
    # palette.setColor(QPalette.ColorRole.HighlightedText, Qt.GlobalColor.black)
    
    # app.setPalette(palette)
    
    window = BlogGeneratorApp()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()